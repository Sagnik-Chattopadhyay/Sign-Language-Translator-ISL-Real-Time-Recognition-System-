from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

MARKDOWN_FILE = r'C:/Users/Sagnik/.gemini/antigravity/brain/42e643a9-9dff-4014-ab5a-82bdbacfeba7/architecture_docs.md'
OUTPUT_FILE = r'e:/projects/sign language/advanced_model/architecture_docs.docx'

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Title
    doc.add_heading('Sign Language Translator Architecture', 0)

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print(f"Error reading markdown: {e}")
        return

    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()
        
        # Code Block
        if stripped.startswith('```'):
            if in_code_block:
                # End of block
                in_code_block = False
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 100) # Dark Blue
                p.style = 'No Spacing' # Or Normal
                code_buffer = []
            else:
                # Start of block
                in_code_block = True
            continue
            
        if in_code_block:
            code_buffer.append(line.rstrip())
            continue

        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        
        # Lists
        elif stripped.startswith('* ') or stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped[0:2].isdigit() and stripped[2:].startswith('. '):
             doc.add_paragraph(stripped.split('. ', 1)[1], style='List Number')

        # Math (Simplified)
        elif '$$' in stripped:
             p = doc.add_paragraph()
             run = p.add_run(stripped.replace('$$', ''))
             run.italic = True
             run.font.color.rgb = RGBColor(100, 0, 100) # Purple

        # Text
        elif stripped:
            doc.add_paragraph(stripped)

    try:
        doc.save(docx_path)
        print(f"Successfully created: {docx_path}")
    except Exception as e:
        print(f"Error saving DOCX: {e}")

if __name__ == "__main__":
    convert_md_to_docx(MARKDOWN_FILE, OUTPUT_FILE)
